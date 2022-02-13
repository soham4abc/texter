import cv2
import os, argparse
import pytesseract
from PIL import Image
import docx
import aspose.words as aw

# We then Construct an Argument Parser


def image_to_text(path):
    # We then read the image with text
    images = cv2.imread(path)

    # convert to grayscale image
    gray = cv2.cvtColor(images, cv2.COLOR_BGR2GRAY)

    # memory usage with image i.e. adding image to memory
    filename = "{}.png".format(os.getpid())
    cv2.imwrite(filename, gray)
    text = pytesseract.image_to_string(Image.open(filename))
    os.remove(filename)
    # print(text)
    # doc = aw.Document()
    # builder = aw.DocumentBuilder(doc)
    # builder.write(text)
    # doc.save("../../Frontend/documents/MyFile.docx")
    finaltext = "".join(c for c in text if valid_xml_char_ordinal(c))
    mydoc = docx.Document()
    # print(type(text))
    mydoc.add_paragraph(finaltext)
    mydoc.save("../../Frontend/documents/MyFile.docx")
    # doc.save("../../Frontend/documents/MyFile.docx")
    # file1 = open("../../Frontend/documents/MyFile.txt", "w")
    # file1.write(text)

    # show the output images
    # cv2.imshow("Image Input", images)
    # cv2.imshow("Output In Grayscale", gray)
    cv2.waitKey(0)


def valid_xml_char_ordinal(c):
    codepoint = ord(c)
    # conditions ordered by presumed frequency
    return (
        0x20 <= codepoint <= 0xD7FF
        or codepoint in (0x9, 0xA, 0xD)
        or 0xE000 <= codepoint <= 0xFFFD
        or 0x10000 <= codepoint <= 0x10FFFF
    )
